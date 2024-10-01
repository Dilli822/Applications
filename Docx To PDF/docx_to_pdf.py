from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
import os

def convert_docx_to_pdf(docx_path):
    # Check if the input file exists
    if not os.path.isfile(docx_path):
        print(f"Error: The file '{docx_path}' does not exist.")
        return

    # Generate output filename based on input filename
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    output_path = os.path.join(os.path.dirname(docx_path), f'{base_name}.pdf')  # Create output filename

    try:
        # Create a PDF document
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        elements = []

        # Load the DOCX document
        doc = Document(docx_path)
        styles = getSampleStyleSheet()

        # Iterate through paragraphs in the DOCX file
        for paragraph in doc.paragraphs:
            # Add text to PDF with styles
            if paragraph.style.name.startswith('Heading'):
                # Use a larger font for headings
                elements.append(Paragraph(paragraph.text, styles['Heading1']))
            else:
                elements.append(Paragraph(paragraph.text, styles['Normal']))
            elements.append(Spacer(1, 0.1 * inch))  # Add space between paragraphs

        # Iterate through tables in the DOCX file
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
            pdf_table = Table(data)
            pdf_table.setStyle(TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                            ('GRID', (0, 0), (-1, -1), 1, colors.black)]))
            elements.append(pdf_table)
            elements.append(Spacer(1, 0.1 * inch))  # Add space after the table

        # Iterate through images in the DOCX file
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image = rel.target_part.blob
                image_path = f'temp_image.png'
                with open(image_path, 'wb') as img_file:
                    img_file.write(image)
                img = Image(image_path)
                img.drawHeight = 1.5 * inch
                img.drawWidth = 1.5 * inch
                elements.append(img)
                elements.append(Spacer(1, 0.1 * inch))  # Add space after the image

        # Build the PDF
        pdf.build(elements)
        print(f"Successfully converted '{docx_path}' to '{output_path}'.")
    except Exception as e:
        print(f"Error converting '{docx_path}': {e}")

# Usage
convert_docx_to_pdf('ML Resume= Dilli Hang Rai.docx')  # Ensure the file has a .docx extension
